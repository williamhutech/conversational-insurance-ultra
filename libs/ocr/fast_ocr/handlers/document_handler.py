"""
Document handler for Word, OpenDocument, RTF, and EPUB files.

Extracts text from various document formats.
"""

from pathlib import Path
from typing import Dict, Any, List
import zipfile
import re

from .base import BaseFileHandler
from ..config import ExtractionConfig
from ..utils.text_utils import clean_ocr_text, merge_text_blocks


class DocumentHandler(BaseFileHandler):
    """
    Handler for document files.

    Supported formats:
    - Microsoft Word: .docx, .dotx
    - Legacy Word: .doc (limited support)
    - OpenDocument: .odt
    - Rich Text Format: .rtf
    - EPUB: .epub
    - Email: .msg, .eml
    """

    @property
    def supported_extensions(self) -> List[str]:
        """Return list of supported document extensions."""
        return [
            '.docx', '.dotx', '.doc', '.odt', '.rtf', '.epub',
            '.msg', '.eml'
        ]

    @property
    def requires_ocr(self) -> bool:
        """Documents don't require OCR."""
        return False

    def extract_text(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """
        Extract text from document file.

        Args:
            file_path: Path to document file
            config: Extraction configuration

        Returns:
            Dictionary with extracted text and metadata
        """
        extension = file_path.suffix.lower()

        try:
            if extension in ['.docx', '.dotx']:
                return self._extract_docx(file_path, config)
            elif extension == '.odt':
                return self._extract_odt(file_path, config)
            elif extension == '.rtf':
                return self._extract_rtf(file_path, config)
            elif extension == '.epub':
                return self._extract_epub(file_path, config)
            elif extension in ['.eml', '.msg']:
                return self._extract_email(file_path, config)
            elif extension == '.doc':
                return self._extract_doc_legacy(file_path, config)
            else:
                raise ValueError(f"Unsupported document format: {extension}")

        except Exception as e:
            return self._create_result(
                text="",
                page_count=0,
                confidence=0.0,
                metadata={'error': str(e), 'extension': extension}
            )

    def _extract_docx(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            return self._create_result(
                text="",
                page_count=0,
                confidence=0.0,
                metadata={'error': 'python-docx not installed'}
            )

        doc = Document(str(file_path))
        paragraphs = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        text = '\n\n'.join(paragraphs)

        if config.enable_text_cleanup:
            text = clean_ocr_text(text)

        return self._create_result(
            text=text,
            page_count=1,  # DOCX doesn't have explicit pages
            confidence=1.0,
            metadata={
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
        )

    def _extract_odt(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Extract text from OpenDocument Text file."""
        try:
            from odf import text, teletype
            from odf.opendocument import load
        except ImportError:
            return self._create_result(
                text="",
                page_count=0,
                confidence=0.0,
                metadata={'error': 'odfpy not installed'}
            )

        doc = load(str(file_path))
        paragraphs = []

        # Extract all text paragraphs
        for para in doc.getElementsByType(text.P):
            para_text = teletype.extractText(para)
            if para_text.strip():
                paragraphs.append(para_text)

        extracted_text = '\n\n'.join(paragraphs)

        if config.enable_text_cleanup:
            extracted_text = clean_ocr_text(extracted_text)

        return self._create_result(
            text=extracted_text,
            page_count=1,
            confidence=1.0,
            metadata={'paragraph_count': len(paragraphs)}
        )

    def _extract_rtf(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Extract text from RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            return self._create_result(
                text="",
                page_count=0,
                confidence=0.0,
                metadata={'error': 'striprtf not installed'}
            )

        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_content = f.read()

        text = rtf_to_text(rtf_content)

        if config.enable_text_cleanup:
            text = clean_ocr_text(text)

        return self._create_result(
            text=text,
            page_count=1,
            confidence=1.0,
            metadata={'format': 'rtf'}
        )

    def _extract_epub(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Extract text from EPUB file."""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback to manual ZIP extraction
            return self._extract_epub_manual(file_path, config)

        book = epub.read_epub(str(file_path))
        chapters = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Parse HTML content
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                text = soup.get_text()
                if text.strip():
                    chapters.append(text)

        text = merge_text_blocks(chapters, separator='\n\n---\n\n')

        if config.enable_text_cleanup:
            text = clean_ocr_text(text)

        return self._create_result(
            text=text,
            page_count=len(chapters),
            confidence=1.0,
            metadata={'chapter_count': len(chapters), 'format': 'epub'}
        )

    def _extract_epub_manual(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Fallback EPUB extraction without ebooklib."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return self._create_result(
                text="",
                page_count=0,
                confidence=0.0,
                metadata={'error': 'BeautifulSoup4 not installed'}
            )

        chapters = []

        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            for file_name in zip_ref.namelist():
                if file_name.endswith(('.html', '.xhtml', '.htm')):
                    with zip_ref.open(file_name) as f:
                        soup = BeautifulSoup(f.read(), 'html.parser')
                        text = soup.get_text()
                        if text.strip():
                            chapters.append(text)

        text = merge_text_blocks(chapters, separator='\n\n---\n\n')

        if config.enable_text_cleanup:
            text = clean_ocr_text(text)

        return self._create_result(
            text=text,
            page_count=len(chapters),
            confidence=0.9,  # Slightly lower confidence for fallback method
            metadata={'chapter_count': len(chapters), 'format': 'epub', 'method': 'manual'}
        )

    def _extract_email(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Extract text from email file (.eml, .msg)."""
        import email
        from email import policy

        with open(file_path, 'rb') as f:
            msg = email.message_from_binary_file(f, policy=policy.default)

        # Extract headers
        headers = []
        for header in ['From', 'To', 'Subject', 'Date']:
            value = msg.get(header)
            if value:
                headers.append(f"{header}: {value}")

        # Extract body
        body_parts = []
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    body_parts.append(part.get_content())
        else:
            if msg.get_content_type() == 'text/plain':
                body_parts.append(msg.get_content())

        # Combine headers and body
        text = '\n'.join(headers) + '\n\n' + '\n\n'.join(body_parts)

        return self._create_result(
            text=text,
            page_count=1,
            confidence=1.0,
            metadata={'format': 'email', 'subject': msg.get('Subject', '')}
        )

    def _extract_doc_legacy(self, file_path: Path, config: ExtractionConfig) -> Dict[str, Any]:
        """Limited support for legacy .doc files."""
        return self._create_result(
            text="",
            page_count=0,
            confidence=0.0,
            metadata={
                'error': 'Legacy .doc format not fully supported. Please convert to .docx',
                'suggestion': 'Use LibreOffice or MS Word to convert .doc to .docx'
            }
        )
